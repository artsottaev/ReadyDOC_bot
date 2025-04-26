from docx import Document

# Функция для генерации .docx документа
async def generate_docx(document_text: str):
    doc = Document()
    doc.add_heading('Юридический документ', 0)
    doc.add_paragraph(document_text)

    # Сохранение документа во временный файл
    temp_file_path = "/tmp/document.docx"
    doc.save(temp_file_path)
    
    return temp_file_path
